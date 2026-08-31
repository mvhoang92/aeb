#!/usr/bin/env python3
"""Export report_v3.md to DOCX using the established report styles."""

from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


REPORT_DIR = Path(__file__).resolve().parent
MARKDOWN = REPORT_DIR / "report_v3.md"
TEMPLATE = REPORT_DIR / "templates" / "aeb_report_template_v3.docx"
OUTPUT_DIR = REPORT_DIR / "exports"
DOCX_PATH = OUTPUT_DIR / "aeb_report_v3.docx"


def clear_document(document):
    body = document._element.body  # pylint: disable=protected-access
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()  # pylint: disable=protected-access
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_repeat_table_header(row):
    properties = row._tr.get_or_add_trPr()  # pylint: disable=protected-access
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, text, separate, end))  # pylint: disable=protected-access


def add_inline(paragraph, text):
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`|(?<!\*)\*[^*]+\*)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def clean_table_cell(value):
    value = value.strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|"):
        value = value[:-1]
    return [part.strip() for part in value.split("|")]


def is_separator_row(cells):
    return bool(cells) and all(re.match(r"^:?-{3,}:?$", cell) for cell in cells)


def add_markdown_table(document, block):
    rows = [clean_table_cell(line) for line in block]
    if len(rows) > 1 and is_separator_row(rows[1]):
        rows.pop(1)
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, values in enumerate(rows):
        for column_index in range(columns):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_inline(paragraph, values[column_index] if column_index < len(values) else "")
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10.5)
                if row_index == 0:
                    run.bold = True
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")
        if row_index == 0:
            set_repeat_table_header(table.rows[row_index])
    document.add_paragraph()


def compatible_image(path, cache_dir):
    suffix = path.suffix.lower()
    if suffix == ".png":
        return path
    cache_dir.mkdir(parents=True, exist_ok=True)
    output = cache_dir / (path.stem + ".png")
    if suffix == ".svg":
        import cairosvg

        cairosvg.svg2png(url=str(path), write_to=str(output), output_width=1800)
    else:
        from PIL import Image

        with Image.open(str(path)) as image:
            image.convert("RGB").save(str(output), "PNG")
    return output


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Trang ")
    add_field(paragraph, "PAGE")
    paragraph.add_run(" / ")
    add_field(paragraph, "NUMPAGES")


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(4)
    for name, size in (("Title", 18), ("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 13)):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
    styles["Caption"].font.name = "Times New Roman"
    styles["Caption"].font.size = Pt(11)


def export():
    document = Document(str(TEMPLATE)) if TEMPLATE.exists() else Document()
    clear_document(document)
    configure_styles(document)
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.5)
    add_page_number(section)

    lines = MARKDOWN.read_text(encoding="utf-8").splitlines()
    image_cache = OUTPUT_DIR / ".report_v3_image_cache"
    index = 0
    first_heading = True
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("```"):
            language = stripped[3:].strip()
            code = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index])
                index += 1
            paragraph = document.add_paragraph()
            paragraph.style = document.styles["Normal"]
            set_cell_shading_placeholder = OxmlElement("w:shd")
            set_cell_shading_placeholder.set(qn("w:fill"), "F2F2F2")
            paragraph._p.get_or_add_pPr().append(set_cell_shading_placeholder)  # pylint: disable=protected-access
            if language:
                paragraph.add_run("[{}]\n".format(language)).bold = True
            run = paragraph.add_run("\n".join(code))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            index += 1
            continue
        if stripped == "$$":
            equation = []
            index += 1
            while index < len(lines) and lines[index].strip() != "$$":
                equation.append(lines[index].strip())
                index += 1
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(" ".join(equation))
            run.font.name = "Cambria Math"
            run.font.size = Pt(12)
            index += 1
            continue
        image = re.match(r"^!\[([^]]*)\]\(([^)]+)\)$", stripped)
        if image:
            image_path = (REPORT_DIR / image.group(2)).resolve()
            if image_path.exists():
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(
                    str(compatible_image(image_path, image_cache)),
                    width=Cm(15.0),
                )
            else:
                document.add_paragraph("[Thiếu hình: {}]".format(image.group(2)))
            index += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            block = []
            while index < len(lines):
                candidate = lines[index].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                block.append(candidate)
                index += 1
            add_markdown_table(document, block)
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if first_heading:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(paragraph, text)
                first_heading = False
            else:
                if level == 1 and (
                    text.startswith("Chương")
                    or text.startswith("Tài Liệu")
                    or text.startswith("Phụ Lục")
                ):
                    document.add_page_break()
                paragraph = document.add_heading(level=min(level, 3))
                add_inline(paragraph, text)
                if text == "Mục Lục":
                    toc = document.add_paragraph()
                    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
            index += 1
            continue
        if re.match(r"^[-*]\s+", stripped):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.7)
            paragraph.paragraph_format.first_line_indent = Cm(-0.4)
            paragraph.add_run("• ")
            add_inline(paragraph, re.sub(r"^[-*]\s+", "", stripped))
            index += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            match = re.match(r"^(\d+)\.\s+", stripped)
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.7)
            paragraph.paragraph_format.first_line_indent = Cm(-0.4)
            paragraph.add_run("{}. ".format(match.group(1)))
            add_inline(paragraph, stripped[match.end() :])
            index += 1
            continue
        if stripped.startswith(">"):
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip().lstrip("> "))
                index += 1
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.7)
            run = paragraph.add_run(" ".join(quote_lines))
            run.italic = True
            continue
        if re.match(r"^\*\*(Hình|Bảng)\s", stripped):
            paragraph = document.add_paragraph(style="Caption")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(paragraph, stripped)
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate or candidate.startswith(("#", "|", "```", "!", ">", "- ", "* ")) or candidate == "$$" or re.match(r"^\d+\.\s+", candidate):
                break
            paragraph_lines.append(candidate)
            index += 1
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(paragraph, " ".join(paragraph_lines))

    document.core_properties.title = "AEB Camera-Radar CARLA Report v3"
    document.core_properties.author = "Mai Việt Hoàng"
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(str(DOCX_PATH))
    cache_dir = OUTPUT_DIR / ".report_v3_image_cache"
    if cache_dir.exists():
        shutil.rmtree(str(cache_dir))
    print("Wrote {}".format(DOCX_PATH))

    with tempfile.TemporaryDirectory(prefix="lo_report_v3_") as profile:
        subprocess.check_call(
            [
                "libreoffice",
                "-env:UserInstallation=file://{}".format(profile),
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(OUTPUT_DIR),
                str(DOCX_PATH),
            ]
        )
    print("Wrote {}".format(OUTPUT_DIR / "aeb_report_v3.pdf"))


if __name__ == "__main__":
    export()
